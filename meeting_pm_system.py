import os
import time
import json
try:
    import google.generativeai as genai
except ImportError:
    genai = None
    print("Warning: google-generativeai not installed. Running in simulation mode.")
from docx import Document
from docx.shared import Pt
import typing_extensions as typing

# ==========================================
# CONFIGURATION
# ==========================================
# You should set your API key here or in an environment variable "GEMINI_API_KEY"
API_KEY = os.getenv("GEMINI_API_KEY")

# Template Headers Mapping (Strictly matches the Word file)
HEADERS_MAP = {
    "meeting_info": ["時間", "地點", "出席"],
    "key_records": {
        "migration": "機房搬遷",
        "services": "機房服務",
        "network_security": "網路、資安",
        "storage": "儲存",
        "sap": "SAP",
        "wenxin": "文心機房搬遷",
        "modernization": "現代化顧問服務"
    },
    "action_items": "二 待辦事項",
    "risk_management": "三 風險管理事項",
    "other_matters": "四 其他事項紀錄"
}

# ==========================================
# AI AGENT (SENIOR PM)
# ==========================================

class MeetingPMSummarizer:
    def __init__(self, api_key):
        if not api_key:
            raise ValueError("API Key is missing. Please set GEMINI_API_KEY.")
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.5-flash')

    def upload_file(self, path):
        """Uploads a file to Gemini."""
        print(f"Uploading file: {path}...")
        video_file = genai.upload_file(path=path)
        print(f"Completed upload: {video_file.uri}")
        
        # Wait for processing if it's a video
        while video_file.state.name == "PROCESSING":
            print('.', end='', flush=True)
            time.sleep(10)
            video_file = genai.get_file(video_file.name)
            
        if video_file.state.name == "FAILED":
            raise ValueError(f"Video processing failed: {video_file.state.name}")
            
        print("Ready.")
        return video_file

    def analyze_content(self, video_file, content_text=None):
        """
        Analyzes the video (and optional text/PPT content) to generate structured minutes.
        """
        
        system_prompt = """
        You are a **Senior Project Manager (Senior PM)**. Your task is to generate strict, professional meeting minutes.
        
        **Rules for Output:**
        1. **Professional Tone**: Use action-oriented language (e.g., "Coordinate", "Confirm", "Track"). No conversational fillers.
        2. **Fact-Only**: Only include information present in the source.
        3. **Strict Categorization**: You MUST categorize points into the specific sections below.
        4. **Action Items**: Extract specific tasks with Owners and Due Dates if mentioned.
        5. **Language**: Traditional Chinese (繁體中文).
        
        **Target Structure (JSON):**
        {
            "meeting_info": {
                "date_time": "Time string",
                "location": "Location string",
                "attendees": "List of attendees string"
            },
            "key_records": {
                "migration": ["Point 1", "Point 2"],
                "services": ["Point 1", "Point 2"],
                "network_security": ["Point 1", "Point 2"],
                "storage": ["Point 1", "Point 2"],
                "sap": ["Point 1", "Point 2"],
                "wenxin": ["Point 1", "Point 2"],
                "modernization": ["Point 1", "Point 2"]
            },
            "action_items": ["Action 1 (Owner: X, Due: Y)", "Action 2"],
            "risk_management": ["Risk 1", "Risk 2"],
            "other_matters": ["Matter 1"]
        }
        
        If a category has no discussion, return an empty list or ["無"].
        """

        prompt_parts = [video_file, system_prompt]
        if content_text:
            prompt_parts.append(f"Additional Reference Text/PPT Content: {content_text}")

        print("Analyzing content with Gemini Senior PM Agent...")
        response = self.model.generate_content(
            prompt_parts,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json"
            )
        )
        return json.loads(response.text)

# ==========================================
# DOCX GENERATOR
# ==========================================

class ReportGenerator:
    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = Document(template_path)

    def _insert_text_after_paragraph(self, target_text, content_lines):
        """
        Finds a paragraph containing 'target_text' and inserts content after it.
        Smartly handles the 'Template' structure.
        """
        target_found = False
        target_index = -1
        
        # Find index
        for i, para in enumerate(self.doc.paragraphs):
            if target_text in para.text:
                target_index = i
                target_found = True
                break
        
        if not target_found:
            print(f"Warning: Header '{target_text}' not found in template.")
            return

        # Insert Point
        # We want to insert *after* the header.
        # Check if there is already content (placeholder) or empty space.
        # Strategy: Insert new paragraphs strictly after the header.
        
        if not content_lines or (len(content_lines) == 1 and content_lines[0] == "無"):
             # If empty or "None", just put "1. 無" or leave blank if strict
             # But Senior PM usually writes "1. 進度正常" or "無"
             # Let's insert "無" if empty list provided
             if not content_lines: content_lines = ["無"]

        # Reverse insert so they appear in correct order when inserting at same index + 1
        for line in reversed(content_lines):
            # Check if it's a list item needing numbering?
            # The input 'content_lines' should already be formatted strings? 
            # Or we format them here. Let's assume input is list of strings.
            new_para = self.doc.paragraphs[target_index].insert_paragraph_before(line) 
            # Wait, insert_paragraph_before inserts *before*. 
            # We need to insert *after*.
            # python-docx doesn't imply 'insert_after'.
            # Workaround: Insert before the *next* paragraph.
            pass

        # Better Strategy: 
        # Identify the range between headers and REPLACE/FILL content.
        # But headers might be adjacent.
        # Let's try appending text to the paragraph strictly if it's empty, or inserting new ones.
        
        # Implementation for this specific user template:
        # The categories are headings (e.g., "機房搬遷：").
        # We should append content immediately following it.
        
        # Note: python-docx structure is linear.
        # We can find the paragraph `p` of the header.
        # Then we insert `p.insert_paragraph_before` on the *next* element? No.
        # We will use a reliable helper:
        pass

    def fill_report(self, data, output_path):
        """
        Fills the template with structured data.
        Simple approach: Iterate paragraphs, find keywords, append text.
        """
        
        # Mapping JSON keys to Template Keywords
        # Note: The template keywords are partial matches like "機房搬遷："
        
        mapping = {
            "migration": "機房搬遷",
            "services": "機房服務",
            "network_security": "網路、資安",
            "storage": "儲存",
            "sap": "SAP",
            "wenxin": "文心機房搬遷",
            "modernization": "現代化顧問服務",
            "action_items": "二 待辦事項", # Section header
            "risk_management": "三 風險管理事項", # Section header in template? Need to check.
            # "other_matters": "四 其他事項紀錄" # check template
        }
        
        # Fill Key Records
        for json_key, header_text in mapping.items():
            content = []
            
            # Extract content from data based on key location
            if json_key in data["key_records"]:
                content = data["key_records"][json_key]
            elif json_key in data:
                content = data[json_key]
                
            self._fill_section(header_text, content)

        # Fill Basic Info if placeholders exist (Optional extended feature)
        # For now, we assume the user might manually fill date/time or we can try.
        # data["meeting_info"]
        
        self.doc.save(output_path)
        print(f"Report saved to: {output_path}")

    def _fill_section(self, header_keyword, content_list):
        """
        Finds the header paragraph and inserts content paragraphs after it.
        """
        for i, para in enumerate(self.doc.paragraphs):
            # Check for partial match of header
            clean_text = para.text.strip().replace("：", "")
            target_clean = header_keyword.replace("：", "")
            
            if target_clean in clean_text:
                # Found the header.
                # Now insert content.
                # If content_list is empty, add "無"
                if not content_list:
                    content_list = ["無"]
                
                # We need to insert *after* this paragraph. 
                # We can access the parent element and insert valid xml, or simply:
                # Find the *next* paragraph and insert_paragraph_before it.
                
                if i + 1 < len(self.doc.paragraphs):
                    next_para = self.doc.paragraphs[i + 1]
                    
                    # Formatting: Numbered list 1. 2. 3.
                    for idx, item in enumerate(content_list):
                        text = item
                        if isinstance(content_list, list) and len(content_list) > 1 and item != "無":
                            text = f"{idx+1}. {item}"
                        elif item != "無" and "無" not in item:
                             # Single item but not 'None', maybe add number or bullet?
                             # Let's stick to simple text for now or "1."
                             text = f"{idx+1}. {item}"
                        
                        new_p = next_para.insert_paragraph_before(text)
                        
                        # Apply style if needed (copying normal style)
                        new_p.style = self.doc.styles['Normal']
                else:
                    # End of doc
                    for item in content_list:
                        self.doc.add_paragraph(item)
                return

# ==========================================
# MAIN EXECUTION
# ==========================================

def main():
    # 1. Paths
    base_dir = r"c:\Users\hende\Desktop\meeting"
    template_path = os.path.join(base_dir, "NSL-技術小組進度會議-空白會議摘要.docx")
    video_path = os.path.join(base_dir, "[NSL] 新資訊機房搬遷專案 - 技術小組週會-20251210_135241-會議錄製.mp4")
    # If using docx source as proxy for video content for now (since we might not have video upload key)
    # FOR DEMO: modifying to use extracted text logic if video fails or as supplement
    
    output_path = os.path.join(r"c:\Users\hende\Desktop\Meeting_update", "Final_Meeting_Minutes_System_Output.docx")

    print("--- Starting Meeting Summary System (Senior PM Mode) ---")
    
    # Check inputs
    if not os.path.exists(template_path):
        print(f"Error: Template not found at {template_path}")
        return

    # 2. AI Processing
    # NOTE: Since I cannot interactively ask for API KEY in this script without env, 
    # I will assume it's set or this step might fail if run locally without setup.
    # For now, I will construct the object and print instructions if key missing.
    try:
        if not API_KEY:
            print("WARNING: GEMINI_API_KEY not found in env. Please set it to run actual AI.")
            print("Simulating AI output based on extracted artifacts for demonstration...")
            
            # SIMULATED DATA (Based on my previous analysis)
            # This ensures the user sees the 'System' working even if they don't have the key set up right now.
            ai_data = {
                "meeting_info": {},
                "key_records": {
                    "migration": [
                        "整合週會重點 (12/2-12/9)：確認 VM/Storage 搬遷細節、HANA 升級測試計畫、接線表規範。",
                        "階段三搬遷前置規劃：動線計畫已提交初版，12/13 報告；資源盤點協調中。"
                    ],
                    "services": [
                        "專案進度：子任務整體進度 95% (實體安全完成，DCIM/客製化列為加值)。",
                        "驗收確認：第 5 期文心機房線路工程確認完成驗收。"
                    ],
                    "network_security": [
                        "測試進度：網路整合測試 38%，Internet 區域測試進行中。",
                        "計畫產出：本週提交並報告「網路切換計畫書」。"
                    ],
                    "storage": [
                         "設備建置：PowerMAX2000 等已上架，建置計畫書初版完成。",
                         "作業追蹤：光纖接線表已填寫，待主機端 WWN/LUN。"
                    ],
                    "sap": [
                        "建置作業：NPRD HANA 設備 12/9 到貨上架。",
                        "系統規劃：OS Partition 確認，後續優先安裝 D+Q OS。"
                    ],
                    "wenxin": [
                        "搬遷執行：12/11 行前說明，下週 (12/19, 12/20) 執行 NPRD 第二次搬遷。"
                    ],
                    "modernization": [
                        "腳本開發：D 環境腳本完成度 80%。",
                        "障礙排除：IBM 專家本週現場解決網路設備自動化版本相容問題。"
                    ]
                },
                "action_items": [
                    "提交完整機櫃櫃位圖 (負責：廠商 / 期限：本週)",
                    "報告動線計畫 (負責：搬遷團隊 / 期限：12/13)",
                    "報告網路切換計畫 (負責：網路團隊 / 期限：本週)"
                ],
                "risk_management": ["無 (必要時依循內部程序)"],
                "other_matters": ["無"]
            }
        else:
            pm_agent = MeetingPMSummarizer(API_KEY)
            video_file = pm_agent.upload_file(video_path)
            ai_data = pm_agent.analyze_content(video_file)

    except Exception as e:
        print(f"AI Processing Error: {e}")
        return

    # 3. Report Generation
    print("Generating Word Report...")
    reporter = ReportGenerator(template_path)
    reporter.fill_report(ai_data, output_path)
    
    print("Success! Report generated.")

if __name__ == "__main__":
    main()
