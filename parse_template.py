# -*- coding: utf-8 -*-
"""
解析 Word 會議記錄範本結構
"""
from docx import Document
from docx.shared import Pt
import os
import sys

# 設定輸出編碼
sys.stdout.reconfigure(encoding='utf-8')

def analyze_document_structure(file_path):
    """分析 Word 文件結構"""
    doc = Document(file_path)
    
    print(f"\n{'='*80}")
    print(f"[FILE] 檔案: {os.path.basename(file_path)}")
    print(f"{'='*80}")
    
    # 分析段落
    print(f"\n[PARAGRAPHS] 段落結構 (共 {len(doc.paragraphs)} 個段落):")
    print("-" * 60)
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            style_name = para.style.name if para.style else "無樣式"
            # 檢查是否為標題
            is_heading = "Heading" in style_name or "標題" in style_name
            
            # 取得字型資訊
            font_info = ""
            if para.runs:
                run = para.runs[0]
                if run.font.size:
                    font_info = f"字級: {run.font.size.pt}pt"
                if run.font.bold:
                    font_info += " [粗體]"
            
            prefix = ">>> " if is_heading else "    "
            text_preview = para.text[:80] + "..." if len(para.text) > 80 else para.text
            print(f"{prefix}[{i}] 樣式: {style_name:20} {font_info}")
            print(f"      內容: {text_preview}")
    
    # 分析表格
    print(f"\n[TABLES] 表格結構 (共 {len(doc.tables)} 個表格):")
    print("-" * 60)
    
    for t_idx, table in enumerate(doc.tables):
        print(f"\n  表格 {t_idx + 1}: {len(table.rows)} 列 x {len(table.columns)} 欄")
        for r_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.replace('\n', ' ')[:30]
                if len(cell.text) > 30:
                    cell_text += "..."
                cells.append(cell_text)
            print(f"    列 {r_idx}: {cells}")
            if r_idx > 8:  # 只顯示前幾列
                print(f"    ... (還有 {len(table.rows) - r_idx - 1} 列)")
                break
    
    return doc

# 分析三個範本
# 分析三個範本
templates = [
    "NSL-技術小組進度會議-20251210會議摘要.docx"
]

base_path = r"c:\Users\hende\OneDrive\桌面\Meeting_update"

for template in templates:
    file_path = os.path.join(base_path, template)
    if os.path.exists(file_path):
        doc = Document(file_path)
        print(f"\n[FULL CONTENT] {template}")
        for para in doc.paragraphs:
            if para.text.strip():
                print(f"Content: {para.text}")
    else:
        print(f"[ERROR] 找不到檔案: {template}")

print("\n\n" + "="*80)
print("[DONE] 分析完成!")
