
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_from_json(json_data, output_path="meeting_minutes.docx"):
    document = Document()
    
    # Setup styles for Traditional Chinese
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    
    # 1. Title
    minutes = json_data.get("meeting_minutes", {})
    title = minutes.get("title", "會議記錄")
    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. Metadata Table
    metadata = minutes.get("metadata", {})
    table = document.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Map metadata to rows
    meta_mapping = [
        ("日期 (Date)", metadata.get("date", "")),
        ("時間 (Time)", metadata.get("time", "")),
        ("地點 (Location)", metadata.get("location", "")),
        ("記錄人 (Recorder)", metadata.get("recorder", "")),
        ("散會時間 (Adjournment)", metadata.get("adjournment_time", ""))
    ]
    
    for i, (label, value) in enumerate(meta_mapping):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        
    document.add_paragraph() # Spacer

    # 3. Attendees
    document.add_heading('出席人員 (Attendees)', level=2)
    attendees = minutes.get("attendees", {})
    
    att_mapping = [
        ("南山高層", attendees.get("nanshan_executives", [])),
        ("技術小組", attendees.get("technical_team_representatives", [])),
        ("PM 代表", attendees.get("pm_representatives", [])),
        ("IBM 代表", attendees.get("ibm_representatives", [])),
        ("參與廠商", attendees.get("participating_vendors", []))
    ]
    
    for category, people in att_mapping:
        p = document.add_paragraph()
        run_label = p.add_run(f"{category}：")
        run_label.bold = True
        p.add_run(", ".join(people))

    # 4. Discussion Records
    document.add_heading('討論事項 (Discussion Records)', level=2)
    discussions = minutes.get("discussion_records", {})
    
    # 4.1 Overall Progress
    section_one = discussions.get("section_one_key_records", {})
    overall = section_one.get("overall_progress_summary", "")
    if overall:
        document.add_heading('整體進度摘要', level=3)
        document.add_paragraph(overall)
        
    # 4.2 Sub Sections
    sub_sections = section_one.get("sub_sections", {})
    sub_mapping = [
        ("機房搬遷 (Server Room Relocation)", "server_room_relocation"),
        ("機房服務 (Server Room Services)", "server_room_services"),
        ("網路與資安 (Network & Security)", "network_and_security"),
        ("儲存 (Storage)", "storage"),
        ("SAP (HW)", "sap_hw"),
        ("文心搬遷 (Wenxin Relocation)", "wenxin_relocation"),
        ("現代化顧問服務 (Modernization Services)", "modernization_services")
    ]
    
    for title, key in sub_mapping:
        item = sub_sections.get(key, {})
        if item:
            h = document.add_heading(title, level=3)
            if "owner" in item:
                p = document.add_paragraph()
                p.add_run("負責人: ").bold = True
                p.add_run(item["owner"])
            
            for detail in item.get("details", []):
                document.add_paragraph(detail, style='List Bullet')

    # 5. Action Items
    document.add_heading('待辦事項 (Action Items)', level=2)
    todos = discussions.get("section_two_todos", [])
    if todos:
        todo_table = document.add_table(rows=1, cols=2)
        todo_table.style = 'Table Grid'
        hdr_cells = todo_table.rows[0].cells
        hdr_cells[0].text = '負責人 (Owner)'
        hdr_cells[1].text = '任務內容 (Task)'
        
        for todo in todos:
            row_cells = todo_table.add_row().cells
            row_cells[0].text = todo.get("owner", "")
            row_cells[1].text = todo.get("task", "")
            
    document.add_paragraph()

    # 6. Risks
    document.add_heading('風險項目 (Risks)', level=2)
    risks = discussions.get("section_three_risks", [])
    if risks:
        for risk in risks:
            p = document.add_paragraph(style='List Number')
            run_title = p.add_run(f"[{risk.get('risk_item', '')}] ")
            run_title.bold = True
            run_title.font.color.rgb = None # Default color or set to red if needed
            p.add_run(risk.get("description", ""))

    # 7. Other Notes
    document.add_heading('其他備註 (Other Notes)', level=2)
    others = discussions.get("section_four_others", [])
    for note in others:
        document.add_paragraph(note, style='List Bullet')

    document.save(output_path)
    print(f"Word document saved to: {output_path}")

if __name__ == "__main__":
    json_data = {
      "meeting_minutes": {
        "title": "南山人壽 114年度新機房基礎架構建置技術小組進度會議紀錄",
        "metadata": {
          "date": "2025/12/30",
          "time": "下午2時00分",
          "location": "Microsoft Teams 線上會議 / 民權東路二段144號7樓會議室",
          "recorder": "PMO",
          "adjournment_time": "下午3時00分"
        },
        "attendees": {
          "nanshan_executives": [
            "Fanny Chan (或代理人)"
          ],
          "technical_team_representatives": [
            "技術小組成員 (依據過往會議名單)",
            "Troy Tsuei",
            "Hank Hsieh",
            "Jason Yu",
            "Jack Sie",
            "Yuan Liu",
            "Erich Lee"
          ],
          "pm_representatives": [
            "KuanLing Lin",
            "WeiYu Chang",
            "Chean Wu"
          ],
          "ibm_representatives": [
            "Qi Peng",
            "Tony Yo",
            "Jack Wang",
            "Eric Chung",
            "Jacqueline Cheng",
            "Jiawei"
          ],
          "participating_vendors": [
            "遠傳",
            "中華",
            "精誠",
            "HPE",
            "Dell",
            "邁達特",
            "仁大"
          ]
        },
        "discussion_records": {
          "section_one_key_records": {
            "overall_progress_summary": "專案整體進度 43%。主任務搬遷 31%，機房服務 95%，網路資安 63%，儲存 8%，SAP升級 29%，文心搬遷 78%，顧問服務 27%。",
            "sub_sections": {
              "server_room_relocation": {
                "owner": "Jack",
                "details": [
                  "本週進度：新機房設備機櫃位置圖（櫃位圖）確認中（70%）；網路 Port Mapping 與 SAN Port Mapping 資料填寫確認中（40%）；平台搬遷方法評估（80%）；搬遷梯次規劃（80%）。",
                  "討論重點：應用系統關聯分析與梯次規劃說明（12/23）；VM 與 Storage 搬遷方法討論（12/29）；應用服務測試計畫文件說明（12/29）。",
                  "下週規劃：持續進行 Port Mapping 與機櫃圖確認；12/31 前確認平台搬遷方法評估初版及搬遷梯次規劃初版；1/31 前完成零件及技術支援計畫規劃。"
                ]
              },
              "server_room_services": {
                "owner": "Jiawei",
                "details": [
                  "本週進度：DCIM 系統安裝、IT 設備建模與建檔進行中；硬體 IO 點位收容彙整；SNMP 連接 Rack controller 測試。",
                  "客製化進度：重新檢討 DCIM 功能清單，列出 99 筆客製化功能待討論；資產管理系統 Prototype 製作及 Review。",
                  "是方機房：確認整合及客製各項工作時程。"
                ]
              },
              "network_and_security": {
                "owner": "Hank & Jason",
                "details": [
                  "網路整合測試進度 61%：Internet 區域 82%，Extranet 區域 75%，Intranet 區域 57%。",
                  "本週重點：聯外線路盤點與測試已於 12/26 完成彙整；持續進行 Internet、Extranet、Intranet 區域整合驗測。",
                  "下週規劃：安排 Server Farm 區域整合測試；12/29~1/2 進行 DMZ 區域整合測試。",
                  "資安：整合測試進度 61%，時程符合預期。"
                ]
              },
              "storage": {
                "owner": "Storage Team",
                "details": [
                  "光纖接線表：Storage 端已填寫完畢，目前等待主機端提供完整的 WWN & LUN 資訊（預計一月份提供）。",
                  "IP/防火牆：新購設備 Management IP 已申請完成，Backup IP 待申請（不影響進度）；防火牆申請剩文心機房部分。",
                  "設備進度：12/23 C30 進貨；12/24 Dell 設備進是方機房上架；12/26 民權 Metro Node 上架；12/30 Dell 設備進文心機房上架。",
                  "下週規劃：持續與 IBM 討論平台搬遷方法；進行 Dell 與 NetApp 設備建置。"
                ]
              },
              "sap_hw": {
                "owner": "Johnny",
                "details": [
                  "機櫃進駐：12/26 提供 HPE 原廠機櫃進駐是方機房時間（1/5 五櫃、1/9 四櫃、1/15 一櫃）。",
                  "DR 環境：CarZ 整併，SDFlex 機箱已定位，待 Upgrade kit 到貨後安裝。",
                  "PRD 環境：進行資安申請前置作業及設備入館申請；硬體設備建置計畫初版產出。",
                  "近期里程碑：1/5 第一批 PRD 設備到貨及上架。"
                ]
              },
              "wenxin_relocation": {
                "owner": "Wenxin Team",
                "details": [
                  "狀態更新：第二次開發測試環境搬遷（Wave 2 - DQ/Sunset），重新安排於 1/16 – 1/17 執行（狀態顯示為延遲結束）。",
                  "本週進度：討論第二梯次搬遷新時程；配合 NPRD Storage 效能測試（台中 Alletra 挖 5 個 1TB LUN 給 QAS 測試）。",
                  "下週規劃：1/5 將 Synergy 設備搬回台中；Unity480 Storage 效能測試；R 環境連線資訊填寫及後視圖整理。"
                ]
              },
              "modernization_services": {
                "owner": "Modernization Team",
                "details": [
                  "腳本進度：S/D 環境測試通過率 81% (56/69)；Q 環境測試通過率 68% (47/69)。",
                  "下週規劃：2026/01 開始建置 AAP (Ansible Automation Platform)，屆時將在 AAP 環境中對腳本進行再次測試與權限管控討論。"
                ]
              }
            }
          },
          "section_two_todos": [
            {
              "owner": "機房搬遷組 (Jack)",
              "task": "12/31 前完成平台搬遷方法評估初版及搬遷梯次規劃初版確認。"
            },
            {
              "owner": "機房搬遷組 (Jack)",
              "task": "1/31 前完成規劃零件及技術支援計畫。"
            },
            {
              "owner": "網路組 (Hank)",
              "task": "12/29~1/2 執行 DMZ 區域整合測試；1/5~1/9 執行 Server Farm 區域整合測試。"
            },
            {
              "owner": "儲存組",
              "task": "追蹤主機端於一月份提供 WWN & LUN 資訊以完成光纖接線表。"
            },
            {
              "owner": "SAP (HW) 組",
              "task": "協助 HPE 原廠於 1/5, 1/9, 1/15 進行機櫃進駐與佈線。"
            },
            {
              "owner": "文心搬遷組",
              "task": "準備於 1/16-1/17 執行延期的第二梯次（Wave 2）搬遷作業。"
            }
          ],
          "section_three_risks": [
            {
              "risk_item": "文心機房搬遷延遲",
              "description": "第二次開發測試環境搬遷（Wave 2）因故重新安排至 1/16-1/17 執行，狀態標示為「延遲結束」，需密切監控是否影響後續梯次。"
            },
            {
              "risk_item": "資訊相依性風險",
              "description": "儲存光纖接線表高度依賴主機端提供 WWN & LUN，目前主機端預計一月份提供，需確保不延誤後續 Storage Zoning 時程。"
            },
            {
              "risk_item": "DCIM 客製化時程",
              "description": "DCIM 提出 99 筆客製化功能，目前整合客製時程仍在討論中，需注意是否影響機房服務驗收。"
            }
          ],
          "section_four_others": [
            "PMO 治理要求：加強 WBS 與時程管理，明確任務相依性（Dependency），並對關鍵節點（Priority -> Super High）進行標註。",
            "目標導向：『以終為始』，確保每一項工作皆為最終交付成果設計，落實風險控管與跨團隊協作。"
          ]
        }
      }
    }
    create_word_from_json(json_data, "c:/Users/hende/Documents/Meeting_update/meeting_minutes_export.docx")
