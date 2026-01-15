
import os
from docx import Document
import sys

# Reconfigure stdout for utf-8
sys.stdout.reconfigure(encoding='utf-8')

base_dir = r"c:\Users\hende\Desktop\meeting"
files = {
    "Source_Candidate_1210": "NSL-技術小組進度會議-20251210會議摘要.docx",
    "Style_Reference_1126": "NSL-技術小組進度會議-20251126會議摘要.docx",
    "Template": "NSL-技術小組進度會議-空白會議摘要.docx"
}

def read_docx(name, filename):
    path = os.path.join(base_dir, filename)
    with open("extracted_content.txt", "a", encoding="utf-8") as f:
        f.write(f"\n{'='*20} {name} ({filename}) {'='*20}\n")
        if not os.path.exists(path):
            f.write("FILE NOT FOUND\n")
            return

        try:
            doc = Document(path)
            f.write(f"Paragraphs: {len(doc.paragraphs)}\n")
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    f.write(f"[{i}] {para.text}\n")
            
            f.write(f"\n--- TABLES ---\n")
            for i, table in enumerate(doc.tables):
                f.write(f"Table {i}:\n")
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    f.write(str(row_text) + "\n")
                    
        except Exception as e:
            f.write(f"Error reading file: {e}\n")

if os.path.exists("extracted_content.txt"):
    os.remove("extracted_content.txt")

for name, filename in files.items():
    read_docx(name, filename)
