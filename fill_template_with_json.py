
import json
import os
from docx import Document
from docx.shared import Pt
import re

# 1. Input Data (Provided by User)
json_data = {
  "meeting_minutes": {
    "title": "南山人壽 114年度新機房基礎架構建置技術小組進度會議紀錄",
    "metadata": {
      "date": "2025/12/30",
      "time": "下午2時00分",
      "location": "Microsoft Teams 線上會議 / 民權東路二段144號7樓會議室",
      "recorder": "PMO",
      "adjournment_time": "下午3時00分"
    },
    "attendees": {
      "nanshan_executives": [
        "Fanny Chan (或代理人)"
      ],
      "technical_team_representatives": [
        "技術小組成員 (依據過往會議名單)",
        "Troy Tsuei",
        "Hank Hsieh",
        "Jason Yu",
        "Jack Sie",
        "Yuan Liu",
        "Erich Lee"
      ],
      "pm_representatives": [
        "KuanLing Lin",
        "WeiYu Chang",
        "Chean Wu"
      ],
      "ibm_representatives": [
        "Qi Peng",
        "Tony Yo",
        "Jack Wang",
        "Eric Chung",
        "Jacqueline Cheng",
        "Jiawei"
      ],
      "participating_vendors": [
        "遠傳",
        "中華",
        "精誠",
        "HPE",
        "Dell",
        "邁達特",
        "仁大"
      ]
    },
    "discussion_records": {
      "section_one_key_records": {
        "overall_progress_summary": "專案整體進度 43%。主任務搬遷 31%，機房服務 95%，網路資安 63%，儲存 8%，SAP升級 29%，文心搬遷 78%，顧問服務 27%。",
        "sub_sections": {
          "server_room_relocation": {
            "owner": "Jack",
            "details": [
              "本週進度：新機房設備機櫃位置圖（櫃位圖）確認中（70%）；網路 Port Mapping 與 SAN Port Mapping 資料填寫確認中（40%）；平台搬遷方法評估（80%）；搬遷梯次規劃（80%）。",
              "討論重點：應用系統關聯分析與梯次規劃說明（12/23）；VM 與 Storage 搬遷方法討論（12/29）；應用服務測試計畫文件說明（12/29）。",
              "下週規劃：持續進行 Port Mapping 與機櫃圖確認；12/31 前確認平台搬遷方法評估初版及搬遷梯次規劃初版；1/31 前完成零件及技術支援計畫規劃。"
            ]
          },
          "server_room_services": {
            "owner": "Jiawei",
            "details": [
              "本週進度：DCIM 系統安裝、IT 設備建模與建檔進行中；硬體 IO 點位收容彙整；SNMP 連接 Rack controller 測試。",
              "客製化進度：重新檢討 DCIM 功能清單，列出 99 筆客製化功能待討論；資產管理系統 Prototype 製作及 Review。",
              "是方機房：確認整合及客製各項工作時程。"
            ]
          },
          "network_and_security": {
            "owner": "Hank & Jason",
            "details": [
              "網路整合測試進度 61%：Internet 區域 82%，Extranet 區域 75%，Intranet 區域 57%。",
              "本週重點：聯外線路盤點與測試已於 12/26 完成彙整；持續進行 Internet、Extranet、Intranet 區域整合驗測。",
              "下週規劃：安排 Server Farm 區域整合測試；12/29~1/2 進行 DMZ 區域整合測試。",
              "資安：整合測試進度 61%，時程符合預期。"
            ]
          },
          "storage": {
            "owner": "Storage Team",
            "details": [
              "光纖接線表：Storage 端已填寫完畢，目前等待主機端提供完整的 WWN & LUN 資訊（預計一月份提供）。",
              "IP/防火牆：新購設備 Management IP 已申請完成，Backup IP 待申請（不影響進度）；防火牆申請剩文心機房部分。",
              "設備進度：12/23 C30 進貨；12/24 Dell 設備進是方機房上架；12/26 民權 Metro Node 上架；12/30 Dell 設備進文心機房上架。",
              "下週規劃：持續與 IBM 討論平台搬遷方法；進行 Dell 與 NetApp 設備建置。"
            ]
          },
          "sap_hw": {
            "owner": "Johnny",
            "details": [
              "機櫃進駐：12/26 提供 HPE 原廠機櫃進駐是方機房時間（1/5 五櫃、1/9 四櫃、1/15 一櫃）。",
              "DR 環境：CarZ 整併，SDFlex 機箱已定位，待 Upgrade kit 到貨後安裝。",
              "PRD 環境：進行資安申請前置作業及設備入館申請；硬體設備建置計畫初版產出。",
              "近期里程碑：1/5 第一批 PRD 設備到貨及上架。"
            ]
          },
          "wenxin_relocation": {
            "owner": "Wenxin Team",
            "details": [
              "狀態更新：第二次開發測試環境搬遷（Wave 2 - DQ/Sunset），重新安排於 1/16 – 1/17 執行（狀態顯示為延遲結束）。",
              "本週進度：討論第二梯次搬遷新時程；配合 NPRD Storage 效能測試（台中 Alletra 挖 5 個 1TB LUN 給 QAS 測試）。",
              "下週規劃：1/5 將 Synergy 設備搬回台中；Unity480 Storage 效能測試；R 環境連線資訊填寫及後視圖整理。"
            ]
          },
          "modernization_services": {
            "owner": "Modernization Team",
            "details": [
              "腳本進度：S/D 環境測試通過率 81% (56/69)；Q 環境測試通過率 68% (47/69)。",
              "下週規劃：2026/01 開始建置 AAP (Ansible Automation Platform)，屆時將在 AAP 環境中對腳本進行再次測試與權限管控討論。"
            ]
          }
        }
      },
      "section_two_todos": [
        {
          "owner": "機房搬遷組 (Jack)",
          "task": "12/31 前完成平台搬遷方法評估初版及搬遷梯次規劃初版確認。"
        },
        {
          "owner": "機房搬遷組 (Jack)",
          "task": "1/31 前完成規劃零件及技術支援計畫。"
        },
        {
          "owner": "網路組 (Hank)",
          "task": "12/29~1/2 執行 DMZ 區域整合測試；1/5~1/9 執行 Server Farm 區域整合測試。"
        },
        {
          "owner": "儲存組",
          "task": "追蹤主機端於一月份提供 WWN & LUN 資訊以完成光纖接線表。"
        },
        {
          "owner": "SAP (HW) 組",
          "task": "協助 HPE 原廠於 1/5, 1/9, 1/15 進行機櫃進駐與佈線。"
        },
        {
          "owner": "文心搬遷組",
          "task": "準備於 1/16-1/17 執行延期的第二梯次（Wave 2）搬遷作業。"
        }
      ],
      "section_three_risks": [
        {
          "risk_item": "文心機房搬遷延遲",
          "description": "第二次開發測試環境搬遷（Wave 2）因故重新安排至 1/16-1/17 執行，狀態標示為「延遲結束」，需密切監控是否影響後續梯次。"
        },
        {
          "risk_item": "資訊相依性風險",
          "description": "儲存光纖接線表高度依賴主機端提供 WWN & LUN，目前主機端預計一月份提供，需確保不延誤後續 Storage Zoning 時程。"
        },
        {
          "risk_item": "DCIM 客製化時程",
          "description": "DCIM 提出 99 筆客製化功能，目前整合客製時程仍在討論中，需注意是否影響機房服務驗收。"
        }
      ],
      "section_four_others": [
        "PMO 治理要求：加強 WBS 與時程管理，明確任務相依性（Dependency），並對關鍵節點（Priority -> Super High）進行標註。",
        "目標導向：『以終為始』，確保每一項工作皆為最終交付成果設計，落實風險控管與跨團隊協作。"
      ]
    }
  }
}

class TemplateReportGenerator:
    def __init__(self, template_path):
        self.doc = Document(template_path)

    def fill_section(self, header_keyword, content_list):
        """
        Finds the paragraph containing header_keyword and inserts content_list after it.
        """
        target_found = False
        target_idx = -1
        
        # 1. Find the header paragraph index
        for i, para in enumerate(self.doc.paragraphs):
            # Normalize text for matching (remove spaces, colons)
            clean_text = re.sub(r"[\s：:\[\]]", "", para.text)
            clean_target = re.sub(r"[\s：:\[\]]", "", header_keyword)
            
            if clean_target in clean_text:
                target_idx = i
                target_found = True
                break
        
        if not target_found:
            print(f"Warning: Header '{header_keyword}' not found in template.")
            return

        # 2. Insert Content
        # We need to insert *after* the header.
        # Since python-docx allows insert_paragraph_before, we target the *next* paragraph.
        # If header is the last paragraph, we append.
        
        if not content_list:
            content_list = ["無"]

        insert_point = None
        if target_idx + 1 < len(self.doc.paragraphs):
            insert_point = self.doc.paragraphs[target_idx + 1]
        
        # Reverse list to keep order correct when using insert_before repeatedly? 
        # No, if we iterate normal order and insert_before the SAME 'insert_point', the order will be reversed in doc.
        # So we should iterate content in REVERSE order if we always insert before the SAME element.
        # OR we update 'insert_point' to be the newly created paragraph? No, that inserts before the new one, reversing again.
        
        # Correct Logic:
        # P_header
        # P_next (insert_point)
        #
        # Action: Insert "Item 1" before P_next.
        # Result:
        # P_header
        # Item 1
        # P_next
        #
        # Action: Insert "Item 2" before P_next.
        # Result:
        # P_header
        # Item 1
        # Item 2
        # P_next
        
        # So standard iteration works perfectly!
        
        for item in content_list:
            text = item
            if insert_point:
                new_p = insert_point.insert_paragraph_before(text)
                new_p.style = self.doc.styles['Normal']
            else:
                self.doc.add_paragraph(text)

    def save(self, path):
        self.doc.save(path)


def main():
    template_path = r"c:\Users\hende\Desktop\meeting\NSL-技術小組進度會議-空白會議摘要.docx"
    output_path = r"c:\Users\hende\Documents\Meeting_update\Final_Meeting_Minutes_Filled.docx"
    
    if not os.path.exists(template_path):
        print(f"Error: Template not found at {template_path}")
        return

    generator = TemplateReportGenerator(template_path)
    
    # --- Prepare Content List from JSON ---
    minutes = json_data["meeting_minutes"]
    discussions = minutes["discussion_records"]
    section_one = discussions["section_one_key_records"]["sub_sections"]
    
    # 1. Main Sections (Sub-sections of Key Records)
    # The template probably has headers like "一、重點紀錄" -> "1. 機房搬遷"
    
    mapping = {
        "機房搬遷": section_one["server_room_relocation"]["details"],
        "機房服務": section_one["server_room_services"]["details"],
        "網路、資安": section_one["network_and_security"]["details"],
        "儲存": section_one["storage"]["details"],
        "SAP": section_one["sap_hw"]["details"], # Keyword might need check, maybe "SAP(HW)"
        "文心機房搬遷": section_one["wenxin_relocation"]["details"],
        "現代化顧問服務": section_one["modernization_services"]["details"],
    }
    
    for header, details in mapping.items():
        # Prepend Owner info to the first detail or as separate line?
        # The JSON has owner info. Let's add it.
        # owner = ... (lookup again)
        # Assuming details is what we want.
        generator.fill_section(header, details)

    # 2. Action Items
    todos = discussions.get("section_two_todos", [])
    todo_lines = []
    for todo in todos:
        todo_lines.append(f"{todo.get('task')} (負責人: {todo.get('owner')})")
    
    generator.fill_section("待辦事項", todo_lines) # Keyword "二、待辦事項" or just "待辦事項"

    # 3. Risks
    risks = discussions.get("section_three_risks", [])
    risk_lines = []
    for risk in risks:
        risk_lines.append(f"[{risk.get('risk_item')}] {risk.get('description')}")
        
    generator.fill_section("風險管理事項", risk_lines)

    # 4. Others
    others = discussions.get("section_four_others", [])
    generator.fill_section("其他事項紀錄", others)

    generator.save(output_path)
    print(f"Document generated at: {output_path}")

if __name__ == "__main__":
    main()
